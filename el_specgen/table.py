import os
from docx.shared import Pt
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document


def create_lists_with_tables(tables_template, lists_template, save_as):
    table = os.getcwd() + tables_template
    doctemplate = Document(table)
    composer = Composer(doctemplate)
    for i in range(0, lists_template - 1):
        doc_temp = Document(table)
        composer.append(doc_temp)
    composer.save(save_as)


def fill_tables_in_template(tables_template):
    temp_doc = Document(tables_template)
    tables = temp_doc.tables
    num = 0
    for j in range(0, len(tables)):
        table = tables[j]
        for i in range(1, 26):
            row = table.rows[i]
            num += 1
            row.cells[0].text = '{{n' + str(num) + '}}'
            row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[1].text = '{{descriptions' + str(num) + '}}'
            row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[2].text = '{{part' + str(num) + '}}'
            row.cells[2].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[3].text = '{{manuf' + str(num) + '}}'
            row.cells[3].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[4].text = '{{unity' + str(num) + '}}'
            row.cells[4].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[5].text = '{{quan' + str(num) + '}}'
            row.cells[5].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)
    temp_doc.save(tables_template)


if __name__ == '__main__':
    create_lists_with_tables("/template/spec_table.docx", 5)
    fill_tables_in_template("template.docx")
