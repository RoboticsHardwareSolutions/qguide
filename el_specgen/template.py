import os
from docx.shared import Pt
from docxtpl import DocxTemplate
from docxcompose.composer import Composer
from docx import Document as Document_compose



def template_merger(templates, tables):
    titul = os.getcwd() + templates.get("template_titul")
    table = os.getcwd() + templates.get("template_table")
    doctemplate = Document_compose(titul)
    composer = Composer(doctemplate)
    for i in range(0, tables):
        doc_temp = Document_compose(table)
        composer.append(doc_temp)
    composer.save('template.docx')
    final_template = DocxTemplate('template.docx')
    return final_template


def template_filler(template):
    tables = template.tables
    num = 0
    for j in range(0, len(tables)):
        table = tables[j]
        for i in range(1, 26):
            row = table.rows[i]
            num += 1
            row.cells[0].text = '{{n' + str(num) + '}}'
            row.cells[0].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[1].text = '{{descriptions' + str(num) + '}}'
            row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[2].text = '{{part' + str(num) + '}}'
            row.cells[2].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[3].text = '{{manuf' + str(num) + '}}'
            row.cells[3].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[4].text = '{{unity' + str(num) + '}}'
            row.cells[4].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)

            row.cells[5].text = '{{quan' + str(num) + '}}'
            row.cells[5].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)
    template.save('template1.docx')



